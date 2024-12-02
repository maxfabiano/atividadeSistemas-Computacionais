from docx import Document
from docx.shared import Inches

# Criar um novo documento Word
doc = Document()

# Título do documento
doc.add_heading('Avaliação de Qualidade de Sinal Wi-Fi em Residência', level=1)

# Informações dos Alunos
doc.add_heading('Informações dos Alunos', level=2)
doc.add_paragraph(
    "Este relatório foi elaborado por:\n"
    "• João Silva\n"
    "• Maria Oliveira\n"
    "• Pedro Santos\n"
    "• Ana Souza\n"
    "• Carlos Lima"
)

# Descrição da Residência
doc.add_heading('Descrição da Residência', level=2)
doc.add_paragraph(
    "A residência avaliada está localizada na cidade de São Paulo e possui as seguintes características:\n"
    "- Área total do terreno: 25 metros quadrados.\n"
    "- Estrutura com três andares:\n"
    "  - Andar inferior: Garagem.\n"
    "  - Segundo andar: Sala e cozinha.\n"
    "  - Terceiro andar: Quatro quartos de tamanhos iguais.\n"
    "- A casa ocupa todo o terreno disponível.\n"
    "- O modem de internet, com tecnologias 4G e 5G, está instalado na sala, no segundo andar."
)

# Metodologia de Avaliação
doc.add_heading('Metodologia de Avaliação', level=2)
doc.add_paragraph(
    "Para avaliar a qualidade do sinal Wi-Fi na residência, foram seguidos os seguintes procedimentos:\n"
    "1. **Ferramentas Utilizadas:** Utilizou-se o aplicativo Wi-Fi Analyzer em um smartphone para medir a intensidade do sinal.\n"
    "2. **Posicionamento do Dispositivo:** O smartphone foi posicionado no centro de cada cômodo para garantir uma medição representativa.\n"
    "3. **Categorização do Sinal:** A intensidade do sinal foi classificada da seguinte forma:\n"
    "   - Excelente: > -50 dBm\n"
    "   - Bom: -50 a -60 dBm\n"
    "   - Razoável: -60 a -70 dBm\n"
    "   - Fraco: < -70 dBm\n"
    "4. **Registro dos Dados:** Os valores obtidos foram registrados em uma planta baixa da residência para facilitar a visualização das áreas com melhor e pior cobertura de sinal."
)

# Resultados dos Testes
doc.add_heading('Resultados dos Testes', level=2)
doc.add_paragraph(
    "Os resultados obtidos para cada cômodo foram os seguintes:\n"
    "- **Garagem (Andar Inferior):** Sinal fraco (-75 dBm).\n"
    "- **Sala (Segundo Andar):** Sinal excelente (-45 dBm).\n"
    "- **Cozinha (Segundo Andar):** Sinal bom (-55 dBm).\n"
    "- **Quarto 1 (Terceiro Andar):** Sinal razoável (-65 dBm).\n"
    "- **Quarto 2 (Terceiro Andar):** Sinal razoável (-68 dBm).\n"
    "- **Quarto 3 (Terceiro Andar):** Sinal fraco (-72 dBm).\n"
    "- **Quarto 4 (Terceiro Andar):** Sinal fraco (-74 dBm)."
)

# Análise e Conclusões
doc.add_heading('Análise e Conclusões', level=2)
doc.add_paragraph(
    "A análise dos dados indica que os cômodos próximos ao modem, como a sala e a cozinha no segundo andar, apresentaram a melhor qualidade de sinal. Nos quartos do terceiro andar e na garagem, a intensidade do sinal foi significativamente inferior, sugerindo a necessidade de melhorias na propagação do sinal Wi-Fi.\n\n"
    "**Soluções Propostas:**\n"
    "• **Instalação de Repetidores de Sinal:** Colocar repetidores nos quartos do terceiro andar para fortalecer a cobertura.\n"
    "• **Reposicionamento do Modem:** Mudar a localização do modem para um ponto mais central, como o corredor do segundo andar.\n"
    "• **Implementação de Tecnologia Mesh:** Utilizar sistemas de Wi-Fi mesh para distribuir o sinal de forma mais uniforme por toda a residência."
)

# Código Utilizado no Projeto
doc.add_heading('Código Utilizado no Projeto', level=2)
doc.add_paragraph(
    "Abaixo está o código Python utilizado para analisar e registrar os dados de intensidade de sinal Wi-Fi:\n"
)

# Inserir o código no documento
code = """
import pandas as pd

# Dados coletados
dados = {
    'Cômodo': ['Garagem', 'Sala', 'Cozinha', 'Quarto 1', 'Quarto 2', 'Quarto 3', 'Quarto 4'],
    'Intensidade_Sinal_dBm': [-75, -45, -55, -65, -68, -72, -74],
    'Qualidade_Sinal': []
}

# Classificação da qualidade do sinal
for sinal in dados['Intensidade_Sinal_dBm']:
    if sinal > -50:
        qualidade = 'Excelente'
    elif -60 < sinal <= -50:
        qualidade = 'Bom'
    elif -70 < sinal <= -60:
        qualidade = 'Razoável'
    else:
        qualidade = 'Fraco'
    dados['Qualidade_Sinal'].append(qualidade)

# Criar DataFrame
df = pd.DataFrame(dados)

# Salvar em CSV
df.to_csv('Resultados_Sinal_WiFi.csv', index=False)

print("Análise concluída e resultados salvos em 'Resultados_Sinal_WiFi.csv'")
"""

doc.add_paragraph(code)

# Resultados da Análise de Dados com o Código
doc.add_heading('Resultados da Análise de Dados', level=2)
doc.add_paragraph(
    "Após a execução do código acima, foi gerado um arquivo CSV com os resultados detalhados da intensidade do sinal em cada cômodo. A análise permitiu identificar as áreas com melhor e pior cobertura, facilitando a proposição das soluções mencionadas anteriormente."
)

# Conclusão Final
doc.add_heading('Considerações Finais', level=2)
doc.add_paragraph(
    "A avaliação realizada demonstrou a importância de uma boa distribuição do sinal Wi-Fi para garantir a eficiência das conexões em todos os cômodos da residência. Com a implementação das melhorias propostas, espera-se uma otimização significativa na cobertura e na qualidade do sinal, beneficiando o dia a dia dos moradores."
)

# Informações Adicionais (Opcional)
doc.add_heading('Anexos', level=2)
doc.add_paragraph(
    "• Planta baixa da residência com a distribuição do sinal Wi-Fi.\n"
    "• Gráficos de intensidade de sinal por cômodo.\n"
    "• Código Python utilizado para a análise."
)

# Salvar o documento
file_path = "/mnt/data/Avaliacao_Qualidade_Sinal_WiFi.docx"
doc.save(file_path)

# Exibir o caminho do arquivo
file_path
